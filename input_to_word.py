from docx import Document
from flask import Flask, request, render_template, current_app
import tkinter.filedialog as filedialog
from tkinter import Tk
from datetime import datetime

app = Flask(__name__)


@app.route("/")
def index():
    # HTMLフォームを表示するためのテンプレートを返す
    return render_template("input_form.html")


@app.route("/get_text", methods=["POST"])
def get_text():
    # フォームから入力されたテキストを取得
    doc_number = request.form["doc_number"]
    doc_date_str = request.form["doc_date"]
    doc_date = datetime.strptime(doc_date_str, "%Y-%m-%d")
    draft_date_str = request.form["draft_date"]
    draft_date = datetime.strptime(draft_date_str, "%Y-%m-%d")
    drafter = request.form["drafter"]
    summary = request.form["summary"]
    authorizere = request.form["authorizere"]
    input_text = request.form["input_text"]

    root = Tk()
    root.withdraw()  # メインウィンドウを非表示にする
    filepath = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])

    # Word文書を生成して保存
    with current_app.app_context():
        generate_docx(filepath, input_text, doc_date, doc_number, draft_date, drafter, summary, authorizere)

    # 応答メッセージを返す
    return "ファイルが保存されました。"


def convert_to_fullwidth_numbers(text):
    # 半角数字と対応する全角数字の辞書を定義
    halfwidth_numbers = "0123456789"
    fullwidth_numbers = "０１２３４５６７８９"
    num_conversion = str.maketrans(halfwidth_numbers, fullwidth_numbers)

    # 変換を行って返す
    return text.translate(num_conversion)


# 全文を改行及び文字数で区切ってリストに入れる
def slice_txt_into_list(full_txt, slice_length):
    sliced_list = []
    lines = full_txt.split("\r\n")  # 改行でテキストを分割

    for line in lines:
        while len(line) > slice_length:
            sliced_list.append(line[:slice_length])
            line = line[slice_length:]

        if line:
            sliced_list.append(line)

    return sliced_list


def generate_docx(output_path, full_txt, date, doc_num, drft_date, drft_person, summary_content, author):
    # 39字ごとに区切ってリストに入れる
    draft_content = slice_txt_into_list(full_txt, 39)

    # 決済者によってWordテンプレートを選択して開く
    if author == "町長":
        document = Document("起案文書様式.docx")
    elif author == "副町長":
        document = Document("起案文書様式(副町長まで).docx")
    else:
        document = Document("起案文書様式(課長まで).docx")

    # 文書内の表を取得する
    table = document.tables[0]

    # 文書番号欄のセルを取得
    doc_num_cell = table.cell(1, 8)

    # 文書番号欄の説に文書番号を入力
    doc_num_cell.text = doc_num

    # 文書の日付欄のセルを取得
    date_cell = table.cell(2, 8)

    # 文書の日付を和暦に変換＆数字を全角に変換
    doc_year_int = int(date.strftime("%Y"))

    doc_year = convert_to_fullwidth_numbers(str(doc_year_int - 2018))
    doc_month = convert_to_fullwidth_numbers(date.strftime("%m"))
    doc_day = convert_to_fullwidth_numbers(date.strftime("%d"))

    # 文書の日付を和暦に変換
    jpn_doc_date = f"令和{doc_year}年{doc_month}月{doc_day}日"

    # 文書の日付欄のセルに日付を入力
    date_cell.text = jpn_doc_date

    # 起案日&起案者欄のセルを取得
    draft_date_cell = table.cell(3, 2)

    # チェック
    # 起案の日付を和暦に変換＆数字を全角に変換
    drft_year_int = int(drft_date.strftime("%Y"))

    drft_year = convert_to_fullwidth_numbers(str(drft_year_int - 2018))
    drft_month = convert_to_fullwidth_numbers(drft_date.strftime("%m"))
    drft_day = convert_to_fullwidth_numbers(drft_date.strftime("%d"))

    # 文書の日付を和暦に変換
    jpn_drft_date = f"令和{drft_year}年{drft_month}月{drft_day}日"

    # 起案日と起案者を同じセルに入力するため連結
    drafter_date = "起案  " + jpn_drft_date + "\n" + "\n" + drft_person

    # 起案日&起案者を入力
    draft_date_cell.text = drafter_date

    # 摘要欄のセルを取得
    summary_content_cell = table.cell(4, 4)

    # 摘要欄に内容を入力
    summary_content_title = "摘要\n" + summary_content
    summary_content_cell.text = summary_content_title

    # 本文を起案文書様式の６行目から入力
    line_num = 6

    for draft in draft_content:
        # 表のline_num行目、0列目にあるセルを取得する
        cell = table.cell(line_num, 0)

        for paragraph in cell.paragraphs:
            paragraph.text = paragraph.text.replace("\r", "")
        # セル内のテキストを上書きする
        cell.text = draft
        line_num += 1

    # 上書きした文書を保存、保存場所指定
    document.save(output_path)


if __name__ == "__main__":
    app.run(host="127.0.0.1", port=8000, debug=True)
